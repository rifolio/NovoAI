import pdfplumber
import pandas as pd
import docx
import os

class PDFProcessor:
    def __init__(self, pdf_path):
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Error: The file {pdf_path} does not exist.")
        
        self.pdf_path = pdf_path
        self.base_name = os.path.splitext(pdf_path)[0]
    
    def extract_data_to_dataframe(self):
        """Extracts tables and text from the PDF and returns a Pandas DataFrame."""
        with pdfplumber.open(self.pdf_path) as pdf:
            tables = []
            extracted_text = []
            
            for page in pdf.pages:
                extracted_table = page.extract_table()
                if extracted_table:
                    tables.extend(extracted_table)

                text = page.extract_text()
                if text:
                    extracted_text.append(text.strip())  # Avoid unnecessary None values

        if not tables:
            raise ValueError("No tables found in the PDF.")

        df = pd.DataFrame(tables)
        
        # Ensure the first row is used as the column header
        df.columns = [f"Column_{i+1}" if not col else col for i, col in enumerate(df.iloc[0])]
        df = df.iloc[1:].reset_index(drop=True)  # Remove duplicate header row

        # Fill empty cells with the value from the cell above (if possible)
        df.fillna(method='ffill', inplace=True)

        return df, extracted_text
    
    def pdf_to_csv(self, output_path=None):
        """Converts extracted table data into a CSV file."""
        df, _ = self.extract_data_to_dataframe()
        output_path = f'{output_path}.csv' if output_path else f"{self.base_name}.csv"
        df.to_csv(output_path, index=False)
        print(f"CSV file saved to {output_path}")
    
    def pdf_to_docx(self, output_path=None):
        """Converts extracted table and text into a Word document."""
        df, extracted_text = self.extract_data_to_dataframe()
        output_path = output_path or f"{self.base_name}.docx"
        
        doc = docx.Document()
        doc.add_heading('Extracted Table from PDF', level=1)

        # Add table to Word
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = "Table Grid"

        # Add column headers
        for col_idx, col_name in enumerate(df.columns):
            table.cell(0, col_idx).text = str(col_name)

        # Add rows
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                table.cell(row_idx + 1, col_idx).text = str(value)

        # Add extracted text
        if extracted_text:
            doc.add_page_break()
            doc.add_heading('Extracted Text from PDF', level=1)
            for text in extracted_text:
                doc.add_paragraph(text)
        
        doc.save(output_path)
        print(f"Word document saved to {output_path}")

# Usage Example
if __name__ == "__main__":
    pdf_path = "PDFiles/FromLaTeX/TableExample1.pdf"
    processor = PDFProcessor(pdf_path)
    processor.pdf_to_csv()
    processor.pdf_to_docx()
